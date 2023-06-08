#===============================================================================
# https://dnevnik.mos.ru/core/api/attestation_subject_lists?class_unit_id=549586
# !!! Не забудьте в методе __init__ класса dn_Auth подставить свои логин/пароль.
#
# Для работа скрипта необходимы библиотеки:
#
# pip install lxml
# pip install pycurl-7.43.0.5-cp39-cp39-win_amd64.whl
# python -m pip install -U pip setuptools
# python -m pip install grab
# pip install pyquery
# pip install xlsxwriter
# pip install python-docx
#
# !!! Библиотеку pycurl лучше ставить именно из указанного whl-контейнера,
#     который предварительно следует скачать на свой комп.
#     Например, с сайта https://www.lfd.uci.edu/~gohlke/pythonlibs/#pycurl
#     Фрагмент "cp39-cp39" в названии файла означает версию установленного
#     на компьютере Python`а. Если установлена 64-битная версия Python`а,
#     то следует ставить "win_amd64"-версию библиотеки.
#     В данный контейнер уже включен файл curl.exe, необходимый для корректной
#     работы библиотеки.
#
# !!! Скрипт тестировался в версии Python 3.9.x - 3.11.x
#     Вышеназванная библиотека pycurl указанной версии может некорректно
#     работать в других версиях Python - в этом случае установите библиотеку
#     pycurl совместимой версии, скачаной с вышеназванного сайта:
#     https://www.lfd.uci.edu/~gohlke/pythonlibs/#pycurl
#
# !!! В whl-контейнере для 11-ой версии Python`а нет самого файла curl.exe -
#     его можно взять из предыдущих версий.
#
#===============================================================================
#
# Консольное приложение!
# Т.е., данный скрипт запускается из командной строки.
#
# Например, из командной строки любого файлового менеджера или можно открыть
# окно командного процессора cmd:
# - в Windows из меню выбрать "Запустить", ввести cmd - откроется окно.
# - с помощью команд вида ">cd Каталог" перейти в каталог со скриптом
# - запустить скрипт (примеры см. ниже)
#
#===============================================================================

import sys, os, fnmatch, zipfile

from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE

import copy

from grab import Grab

#===============================================================================
# Формирование doc-файлов контрольных листов для каждого выпускника.
#
# !!! Работает со справочником предметов нового блока печати аттестатов
#     "Аттестаты (new)".
#
# Обязательные параметры:
# - путь к папке, куда будут сохранены сформированные doc-файлы контролек.
#
# Необязательные параметры:
# - имя класса или параллель, в этом случае в заданную папку будет сохранен
#   zip-архив с контрольками только для этого класса или zip-архивы с
#   контрольками для классов заданной параллели.
#   Класс задается в формате "9-А", где буква - русская.
#   Параллель задается числом 9 или 11.
#   Если заданы одновременно имя класса и параллель, то будет учтен только
#   первый по порядку дополнительный параметр.
#   Например, att_mark_get_docs jrn 9 9-Б
#   В папку jrn будут сохранены zip-архивы с контрольками для всех 9-х классов.
#
# Примеры запуска скрипта:
#
# roa_att_mark_get_docs.py tmp
# В папке "tmp", которая находится в папке запуска скрипта, будут созданы
# zip-архивы с "контрольками" по всем 9-ым и 11-ым классам школы.
#
# roa_att_mark_get_docs.py tmp 9
# В папке "tmp", которая находится в папке запуска скрипта, будут созданы
# zip-архивы с "контрольками" по всем 9-ым классам школы.
#
# roa_att_mark_get_docs.py C:\Temp "НДО-9 Иванов Иван"
# В папке "C:\Temp" будет создан zip-архив с "контролькой" для надомника
# "НДО-9 Иванов Иван".
#
#===============================================================================

class dn_Auth:

      base = ""; user = ""; password = ""
      web = None
      pid = ""; aid = ""; curr_aid = ""

      def __init__(self, dn="test", timeout=15, conn_tm=10, aid=""):
          self.timeout = timeout
          self.conn_tm = conn_tm
          if dn.lower() != "work":
             self.base = "https://dnevnik-test.mos.ru/"
             self.user = "ignatov-ai"
             self.password = "nxkt97vS!@"
          else:
             self.base = "https://dnevnik.mos.ru/"
             self.user = "ignatov-ai"
             self.password = "nxkt97vS!@"

          self.free()
          self.aid = aid

      def __del__(self):
          self.free()

      def free(self):
          if self.web != None:
             self.web = None
          self.pid = ""

      def set_aid(self, aid=""):
          if not self.web: return("")

          aid = aid.strip()
          if (not aid) or (aid == "0"):
             aid = self.curr_aid
          elif aid[0] == "-":
             shift = aid[1:]
             aid = self.curr_aid
             if shift.isdigit() and int(shift):
                shift = int(shift)
                self.web.go(self.base + "core/api/academic_years")
                if (self.web.doc.code == 200) and \
                   (len(self.web.doc.json) > shift):
                   cy = False
                   for y in sorted(self.web.doc.json,             \
                                   key=lambda y: y["begin_date"], \
                                   reverse=True):
                       if y["current_year"]:
                          cy = True
                       elif cy:
                          shift -= 1
                          if not shift:
                             aid = str(y["id"])
                             break

          self.aid = aid
          self.web.config['cookies']['aid'] = self.aid
          self.web.cookies.set('aid', self.aid, \
                               "."+ self.base.split("://")[-1].split("/")[0])

          return(self.aid)

      def login(self):
          if self.web: return True

          # Подготовка к подключению
          self.web = Grab(connect_timeout = self.conn_tm, timeout = self.timeout)
          self.web.config['common_headers']['Accept-Language'] = 'ru-RU'
          self.web.config['common_headers']['Accept'] = 'application/json,'+ \
                         self.web.config['common_headers']['Accept']

          # Запрос страницы входа для получения куков
          self.web.go(self.base)
          if (self.web.doc.code != 200):
             print("Ошибка доступа к сайту \""+ self.base +"\"!!!")
             self.free()
             return False

          # Авторизация
          self.web.setup(headers={'Content-Type': 'application/json;charset=UTF-8'})
          self.web.go(self.base +"lms/api/sessions", \
                      post = '{"login":"%s","password_plain":"%s"}' % \
                             (self.user, self.password),
                      timeout = self.timeout)
          if (self.web.doc.code != 200):
             print("Ошибка авторизации ("+ str(self.web.doc.code) +")!!!")
             self.free()
             return False

          # На всякий случай (несколько учеток) ищем профиль админа
          data = self.web.doc.json
          profile = list(filter(lambda i: "school_admin" in i["roles"], \
                                data["profiles"]))
          if not profile:
             print("Отсутствует профиль админа!")
             self.free()
             return False
          else:
             profile = profile[0]
             self.pid = str(profile["id"])

          # Сохранение куков и данных пользователя
          self.web.setup(connect_timeout=self.conn_tm, timeout=self.timeout)
          self.web.config["cookies"] = { \
                         "auth_token": data["authentication_token"],
                         "profile_id": self.pid,
                         "profile_roles": "%2C".join(profile["roles"]),
                         "is_auth": "true",
                         "user_id": data["id"]}
          self.web.config['common_headers']['Auth-Token'] = \
                         data["authentication_token"]
          self.web.config['common_headers']['Profile-Id'] = self.pid

          # Определение ид текущего учебного года
          start_aid = self.aid
          self.web.go(self.base + "core/api/academic_years")
          if self.web.doc.code == 200:
             aid = list(filter(lambda i: i["current_year"], self.web.doc.json))
             if aid: self.aid = str(aid[0]["id"])
          self.curr_aid = self.aid or "10"
          self.set_aid(start_aid if start_aid else self.curr_aid)
      
          return True

def zipFiles(dirPath, archName="", delFiles=False, fMask="*"):
    # dirPath  - рабочий каталог, файлы в котором подлежат архивации
    # archName - имя архива (может включать полный путь)
    # delFiles - признак удаления архивируемых файлов
    # fMask    - маска фильтра файлов ("*", "?", [список], [!список])
    #            пустая маска - пустой список файлов

    if not len(dirPath) or not os.path.isdir(dirPath):
       return -1

    if not dirPath[-1] in ["\\","/"]: dirPath += "\\"
    files = [f for f in fnmatch.filter(os.listdir(dirPath), fMask) if os.path.isfile(dirPath + f)]
    if not len(files): return -2

    if not len(archName):
       # Если не задано имя архива - совпадает с именем рабочего каталога
       archName = os.path.basename(dirPath[:-1]) +".zip"

    # Если не задан путь к архиву - создаем его в рабочем каталоге
    if not len(os.path.dirname(archName)): archName = dirPath + archName
    try:
       arch = zipfile.ZipFile(archName, "w", zipfile.ZIP_DEFLATED)
    except Exception:
       return -3

    deletedFiles = []
    success = 0
    for f in files:
        fPath = dirPath + f
        if os.path.isfile(fPath):
           try:
              arch.write(fPath, f)
              success += 1
              if delFiles: deletedFiles.append(fPath)
           except Exception:
              success = 0
              break

    arch.close()
    if not success: return -4

    # Если задано - удаляем заархивированные файлы
    for f in deletedFiles:
        try:
           os.remove(f)
        except Exception:
           break

    return success

#===============================================================================
#===============================================================================

# Список аттестационных предметов в порядке вывода в аттестат
subj_names = []
# Словарь по id предмета для формирования аттестационных оценок ученика.
# {subj_id:[order, ""]}
# Для вывода в отчет из этого словаря будет создан массив оценок:
# [i[1][1] for i in sorted(subj_ids.items(),key=lambda i: i[1][0])]
subj_ids = {}

def att_subj_load(_cls_id):
    #================================================================
    # В случае наличия нескольких списков предметов по одному классу
    # предполагаем, что все предметы в правильном порядке включены
    # в список с наибольшим количеством предметов - его и используем.
    #================================================================
    global subj_names, subj_ids

    subj_names = []
    subj_ids = {}

    # Загрузка списка аттестационных предметов заданного класса
    dn.web.go(dn.base +"core/api/attestation_subject_lists"+ \
              "?class_unit_id="+ str(_cls_id))
    if dn.web.doc.code != 200:
       print("- ошибка получения списка аттестационных предметов! ("+ \
             str(dn.web.doc.code) +")")
       return 0
    elif not len(dn.web.doc.json):
       print("- для данного класса не указан список аттестационных предметов!")
       return 0

    subj_list = list(sorted(dn.web.doc.json, \
                            key=lambda s: len(s["subjects"])))[-1]["subjects"]

    order = 0
    for s in sorted(filter(lambda s: \
                    s["attestation_subject_load"] == "REQUIRED", subj_list), \
                    key=lambda s: s["order"]):
        order += 1
        subj_names.append(s["name"])
        subj_ids[s["subject_id"]] = [order, ""]

    for s in sorted(filter(lambda s: \
                    s["attestation_subject_load"] == "OPTIONAL_LOAD_OVER_64", \
                    subj_list), key=lambda s: s["order"]):
        order += 1
        subj_names.append(s["name"])
        subj_ids[s["subject_id"]] = [order+1000, ""]

    for s in sorted(filter(lambda s: \
                    s["attestation_subject_load"] == "OPTIONAL_LOAD_LESS_64", \
                    subj_list), key=lambda s: s["order"]):
        order += 1
        subj_names.append(s["name"])
        subj_ids[s["subject_id"]] = [order+2000, ""]

    if not len(subj_names):
       print("- для данного класса не указан список аттестационных предметов!")

    return len(subj_names)

#===============================================================================

if len(sys.argv) < 2:
   print("Запуск: att_mark_get_docs path_folder [class_name|class_level]")
   print("path_folder - путь к папке для zip-архивов с контрольками по классам")
   print("class_name - имя класса, необязательный параметр")
   print("class_level - параллель, необязательный параметр")
   exit()

path_folder = sys.argv[1].strip()
if path_folder == ".":
   path_folder = ""
elif not os.path.isdir(path_folder):
   print("Указанная папка не существует - задайте путь к существующей папке!")
   exit()
elif path_folder[-1] != "\\":
   path_folder += "\\"

class_level = 0
class_name = ""
if len(sys.argv) > 2:
   class_name = sys.argv[2].strip().upper()
   if (len(class_name) < 3) and class_name.isdigit():
      class_level = int(class_name)
      class_name  = ""
   else:
      class_level = class_name.split("-")[0]
      if class_level.isdigit():
         class_level = int(class_name.split("-")[0])
      else:
         class_level = 0

# Авторизация в ЭЖД
dn = dn_Auth("work", timeout=30)
if (dn is None) or not dn.login():
   print("Ошибка входа в ЭЖД!")
   exit()

# Запрос списка классов
print("Обработка списка классов ЭЖД ...")
if class_level:
   params = "&class_level_id="+ str(class_level)
else:
   params = "&class_level_ids=9,11"

dn.web.go(dn.base +"core/api/class_units?academic_year_id="+ dn.aid + params + \
          "&with_home_based=true")
if dn.web.doc.code != 200:
   print("Ошибка получения списка классов! ("+ str(dn.web.doc.code) +")")
   exit()

months = ["января", "февраля", "марта", "апреля", "мая", "июня", "июля",
          "августа", "сентября", "октября", "ноября", "декабря"]

# Обходим классы
cls_all = 0; cls_saved = 0
for cl in sorted(dn.web.doc.json, key=lambda item: item["home_based"], reverse=True):
    cls_name  = cl.get("name", "").upper()
    if not cls_name: continue
    if class_name and (cls_name != class_name): continue

    cls_id = cl.get("id", 0)
    if not cls_id: continue

    cls_level = cl.get("class_level_id", 0)

    print(cls_name)
    cls_all += 1

    # Запрос списка учеников класса
    dn.web.go(dn.base +"core/api/student_profiles?pid="+ dn.pid +
              "&academic_year_id="+ str(dn.aid) +
              "&class_unit_id="+ str(cls_id) +
              "&with_user_info=true")
    if dn.web.doc.code != 200:
       print("Ошибка получения списка учеников! ("+ str(dn.web.doc.code) +")")
       continue
    elif not len(dn.web.doc.json):
       print("Список учеников пуст!")
       continue

    st_list = dn.web.doc.json
    st_list.sort(key=lambda item: item["user_name"])

    if not att_subj_load(cls_id): continue

    # Формирование списка учеников класса
    cls_st = [] # [st_id, st_name, st_db]

    # Обходим учеников
    for st in st_list:
        st_name = st.get("user_name", "").strip()
        st_id = st.get("id", 0)
        if not st_name or not st_id: continue
        st_bd = st.get("birth_date","")

        # Задана ли дата выбытия
        if st.get("left_on",""): continue
        # Задан ли УП
        if not st.get("curricula",""): continue

        cls_st.append([st_id, st_name, st_bd])

    if not len(cls_st):
       print("Список учеников пуст!")
       continue

    # Запрос аттестационных оценок учеников класса
    marks = {}
    for page in range(1, 1000):
        dn.web.go(dn.base +"core/api/final_marks?mark_type=attestation"+ \
                  "&academic_year_id="+ str(dn.aid) +"&student_profile_ids="+ \
                  ",".join([str(st[0]) for st in cls_st]) + \
                  "&per_page=300&page="+ str(page))
        if dn.web.doc.code != 200:
           print("Ошибка получения аттестационных оценок! ("+ \
                 str(dn.web.doc.code) +")")
           break
        elif (not dn.web.doc.json) or (not len(dn.web.doc.json)):
           break

        # Формирование словаря аттестационных оценок учеников класса
        for subj in dn.web.doc.json:
            subj_id = subj["subject_id"]
            if not subj_id: continue
            st_id = subj["student_profile_id"]
            if not st_id: continue
            mark = subj["value"]

            if not st_id in marks:
               # оценок этого ученика еще нет в словаре - добавляем шаблон
               marks[st_id] = copy.deepcopy(subj_ids)

            if subj_id in marks[st_id]:
               marks[st_id][subj_id][1] = mark

    if dn.web.doc.code != 200: continue

    # Формирование контролек по ученикам класса
    for st in cls_st:
        doc = Document()
        doc.add_heading(st[1], 0)
        bd = st[2].split(".")
        doc.add_paragraph("Дата рождения: "+ str(int(bd[0])) +" "+ \
                          months[int(bd[1])-1] +" "+ bd[2] +" года")
        if cls_level == 9:
           att_name = doc.add_paragraph("Аттестат об основном общем образовании")
        else:
           att_name = doc.add_paragraph("Аттестат о среднем общем образовании")
        doc.add_paragraph('Государственное автономное общеобразовательное учреждение города Москвы Школа №548 "Царицыно"')
        
        st_marks = marks.get(st[0],{})
        if not st_marks:
           doc.add_paragraph("НЕТ ОЦЕНОК!!!")
           gold = False
        else:
           ext_subj = False
           gold = True
           st_marks = [i[1] for i in sorted(st_marks.items(),key=lambda i: i[1][0])]
           table = doc.add_table(rows=0, cols=2)
           table.autofit = True
           table.columns[0].width = Cm(20.0)
           for ndx, mark in enumerate(st_marks):
               if not mark[1]: continue

               if (mark[0] > 2000) and not ext_subj:
                  # Начало секции "Дополнительные сведения"
                  row = table.add_row()
                  row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                  row.height = Cm(0.5)
                  row = table.add_row()
                  row.cells[0].text = "----- Дополнительные сведения -----"
                  row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                  row.height = Cm(0.5)
                  ext_subj = True

               row = table.add_row()
               row.cells[0].text = subj_names[ndx]
               if (type(mark[1]) != str) and (int(mark[1]) > 2):
                  row.cells[1].text = str(int(mark[1]))
                  if int(mark[1]) != 5: gold = False
               row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
               row.height = Cm(0.5)

        if gold:
           att_name.add_run(" (с отличием)").bold = True
        doc.save(path_folder + st[1] +".docx")

    if zipFiles(path_folder, cls_name + " (контрольные листки).zip", True, "*.docx") >= 0:
       cls_saved += 1
    else:
       print("Ошибка создания архива контролек!")
       exit()

print("Создано архивов контролек: "+ str(cls_saved) +" из "+ str(cls_all) +"\n")
